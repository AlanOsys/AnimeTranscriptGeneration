from os import listdir
from os.path import isfile, join
import docx
import docx2txt
import string



f = open("inputAnime.txt", "a", errors="ignore", encoding="utf-8")
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    print(len(doc.paragraphs))
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

onlyfiles = [f for f in listdir('Data') if isfile(join('Data', f))]

data = listdir('Data')
for i in range(len(data)):
    folder = listdir(f'Data/{data[i]}')
    for j in range(len(folder)):
        scripts = listdir(f'Data/{data[i]}/{folder[j]}')
        for k in range(len(scripts)):
            if 'JPN' not in scripts[k] and 'docx' in scripts[k] and 'ROM' not in scripts[k]:
                #text = getText(f'Data/{data[i]}/{folder[j]}/{scripts[k]}')
                text = docx2txt.process(f'Data/{data[i]}/{folder[j]}/{scripts[k]}')
                sep = 'Please read my FAQ for more usage details.'
                stripped = text.split(sep, 1)[1]
                for char in string.punctuation:
                    stripped = stripped.replace(char, ' ')
                f.write(stripped)
f.close()
